from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from docx import Document
from docx.shared import Inches
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
import uvicorn
import os

app = FastAPI()

def last_index(list, value):
    for i in range(len(list) - 1, -1, -1):  # Iterate in reverse
        if list[i] == value:
            return i

class RetainerRequest(BaseModel):
    FirmName: str
    FirmAddress: str
    FirmPhoneNumber: str
    FirmEmail: str
    FirmWebsite: str
    FirmLeadPartner: str
    ClientName: str
    ClientAddress: str
    ContactName: str
    ContactPhoneNumber: str
    ContactEmail: str
    ExistingRelationship: int
    InitialMeetingNotes: Optional[str] = ""
    Scope: str
    BillingType: str
    PartnerRate: str
    AssociateRate: str
    StudentRate: str
    Frequency: str
    RetainerAmount: str
    LateFeePercent: str
    Jurisdiction: str
    DisputeResolutionClause: int
    ConfidentialityClause: int
    FilePath: str = os.path.join(os.environ["USERPROFILE"], "Downloads\\Letter")

@app.post("/generate_retainer")
async def generate_retainer(data: RetainerRequest):
    async def content_generator():
        try:
            tavily_search = TavilySearchResults()
            results = tavily_search.invoke(data.ClientName + " company profile")

            llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)
            first_draft = ""

            prompts = {
                "introduction": f"Write a professional but friendly introduction for a retainer letter from \
                    {data.FirmName} to {data.ClientName}. The tone should be warm, approachable, and reflect a strong client relationship. \
                    If the firm has an existing relationship with the client, acknowledge it. If this is a new engagement, \
                    reference the initial meeting and key points discussed. Use natural, conversational language while maintaining \
                    professionalism. The introduction should allow for a smooth transition into the scope of services the firm will provide \
                    which is a paragraph that will be written later. Write the content only, no need for letter heading or salutation. \
                    Do we have a previous history with the client? {"Yes." if data.ExistingRelationship == 1 else "No."} \
                    Here is some background on the client: {str(results)} \
                    Here is some notes from the initial meeting (may be blank if nothing of interest): {data.InitialMeetingNotes}",
                "services": f"Write a detailed scope of services section for a legal retainer letter between {data.FirmName} and {data.ClientName}. \
                    The section should confirm the engagement, outline the specific legal services being provided, and clarify what is not included. \
                    A scope has been provided below. Clearly state that the firm does not guarantee specific outcomes. Write the content only. \
                    Lawyers other than the lead partner may need to be introduced and a placeholder section should be set up for this. \
                    Include a brief note on client communication policies and document delivery expectations. Use clear, professional, \
                    and legally precise language. The work will most likely be supervised by {data.FirmLeadPartner} in the jurisdiction of {data.Jurisdiction} \
                    Scope: {data.Scope}",
                "fees": f"Write a professional and legally precise fees section for a retainer letter between {data.FirmName} and {data.ClientName}. \
                    The section should clearly explain how fees will be calculated, including whether the arrangement is hourly, fixed-fee, or estimate-based. \
                    If hourly, specify that fees vary by attorney seniority and outline the billing cycle (monthly, semi-monthly, etc.). \
                    If fixed-fee or estimate-based, clarify that the fee covers specific services but may not include disbursements or \
                    unforeseen additional work. Ensure the language is transparent about billing expectations while protecting the firm from \
                    liability for cost overruns. Use professional and client-friendly language. Write the content only. \
                    Billing Type: {data.BillingType} \
                    Partner Rate: {data.PartnerRate} \
                    AssociateRate: {data.AssociateRate} \
                    Articling Student Rate: {data.StudentRate} \
                    Frequency: {data.Frequency} \
                    Retainer Amount: {data.RetainerAmount} \
                    Late Fee Percentage: {data.LateFeePercent} \
                    Please write the content only and use bullet points where applicable.",
                "disbursements": f"Write a professional and legally precise disbursements and other charges section for a retainer letter between \
                    {data.FirmName} and {data.ClientName}. The section should clearly define disbursements as third-party costs incurred on the client's behalf and \
                    distinguish them from office-related charges. List common disbursements relevant to the legal matter, such as government filing fees, \
                    expert witness costs, and travel expenses. If applicable, specify whether disbursements will be billed as incurred or included as a \
                    flat fee. Ensure the language is transparent, protecting the firm from absorbing unexpected costs while setting clear expectations for \
                    the client. Write the content only. Please tailor it to any of the following: {data.Scope} \
                    Initial Meeting Notes: {data.InitialMeetingNotes}",
                "interest": f"Write a clear and legally precise interest on late payments section for a retainer letter between {data.FirmName} and \
                    {data.ClientName}. The section should state when payment is due and specify the interest rate applied to overdue balances. Clarify \
                    whether interest is simple or compound, how it is calculated, and from what date it accrues. Ensure the language is transparent, \
                    reinforcing that clients will not be surprised by interest charges. Maintain a professional yet firm tone to ensure compliance \
                    while preserving client relationships. Write the content only. Interest Rate: {data.LateFeePercent} \
                    Initial Meeting Notes: {data.InitialMeetingNotes}",
                "retainer": f"Write a clear and legally precise financial retainer section for a retainer letter between {data.FirmName} and {data.ClientName}. \
                    The section should specify the retainer amount, explain that it will be held in a trust account, and clarify that fees, disbursements, \
                    and taxes will be deducted from it. Outline the process for replenishment and how any unused funds will be handled upon termination \
                    of services. Maintain a professional and transparent tone to ensure the client understands their financial obligations. \
                    Write the content only. Retainer Amount: {data.RetainerAmount} \
                    Initial Meeting Notes: {data.InitialMeetingNotes}",
                "conclusion": f"Write a clear and professional agreement section for a retainer letter between {data.FirmName} and {data.ClientName}. \
                    The paragraph(s) should request the client to review the letter, ensure they understand its terms, and return a signed copy as confirmation \
                    of engagement. It should mention the importance of seeking independent legal advice if needed and outline the next steps if the client \
                    does not wish to proceed. Maintain a formal yet client-friendly tone. Write the content only. \
                    Scope: {data.Scope} \
                    Initial Meeting Notes: {data.InitialMeetingNotes}"
            }

            if data.DisputeResolutionClause:
                prompts["dispute"] = f"Write a professional and legally precise dispute resolution and termination section for a retainer letter between \
                    {data.FirmName} and {data.ClientName}. The section should outline how concerns about legal services should be addressed, including escalation \
                    to a senior firm member if necessary. It should also explain the client's right to terminate services with written notice and the \
                    law firm's right to withdraw representation under specific circumstances (e.g., non-payment, loss of confidence). Ensure the \
                    language is clear, professional, and protective of both parties. If applicable, mention the firm's process for transferring files \
                    to successor counsel. Write the content only. \
                    Initial Meeting Notes: {data.InitialMeetingNotes}"

            if data.ConfidentialityClause:
                prompts["confidentiality"] = f"Write a professional and legally precise confidentiality clause for a retainer letter between \
                    {data.FirmName} and {data.ClientName}. The section should emphasize that all communications, documents, and information shared between \
                    the firm and the client will be kept strictly confidential, in accordance with legal and ethical obligations. \
                    Clarify any exceptions, such as court orders or legal requirements to disclose certain information. Ensure the language is clear, \
                    legally sound, and protective of both parties. Write the content only. \
                    Initial Meeting Notes: {data.InitialMeetingNotes}"

            for idx, (key, prompt) in enumerate(prompts.items(), start=1):
                yield f"data: Creating {key} section... ({idx}/{len(prompts)})\n"
                first_draft += llm.invoke([HumanMessage(content=prompt)]).content
            
            yield "Writing letter to doc...\n"
            
            review_prompt = f"Please format the following sections into a professional, cohesive, and structured retainer letter. \
                Ensure proper legal tone, grammar, and logical flow. Add appropriate letter formatting, firm and client details, and headers. \
                Improve clarity while maintaining a formal and client-friendly approach. Write the content only. \
                The output must be html. If necessary, adjust the section order but do not lose any sections. \
                Ensure any signature section appears at the end. Here is some background on the client: {str(results)} \
                Firm Details: \
                {data.FirmName} \
                {data.FirmAddress} \
                Lead Partner: {data.FirmLeadPartner} \
                Phone: {data.FirmPhoneNumber} \
                Email: {data.FirmEmail} \
                Website: {data.FirmWebsite} \
                 \
                Client Details: \
                {data.ClientName} \
                {data.ClientAddress} \
                Attention: {data.ContactName} \
                Phone: {data.ContactPhoneNumber} \
                Email: {data.ContactEmail} \
                 \
                Letter: {first_draft}"
            reviewed_draft = llm.invoke([HumanMessage(content=review_prompt)]).content

            doc = Document()
            html = reviewed_draft.split('<')
            for i in range(len(html)):
                if "body>" in html[i]:
                    break
                
            j, style, list_level = i, ["Normal"], 0
            heading_styles = {"h1": "Title", "h2": "Heading 1", "h3": "Heading 2"}
            list_styles = {"ol": "List Number", "ul": "List Bullet"}

            while "/body>" not in html[j]:
                row = html[j].split(">")
                tag, content = row[0], row[1] if len(row) > 1 else ""

                if tag in heading_styles:
                    style.append(heading_styles[tag])
                    doc.add_paragraph(content, style=style[-1])
                elif tag == "p":
                    doc.add_paragraph(content, style=style[-1])
                elif tag in list_styles:
                    style.append(list_styles[tag])
                    list_level += 1
                elif tag == "li" and content.strip():
                    p = doc.add_paragraph(content, style=style[-1])
                    p.paragraph_format.left_indent = Inches(0.5 * list_level)
                elif tag in {"b", "strong", "i"}:
                    p = doc.add_paragraph(style=style[-1])
                    if list_level > 0:
                        p.paragraph_format.left_indent = Inches(0.5 * list_level)
                    run = p.add_run(content)
                    setattr(run, "bold" if tag in {"b", "strong"} else "italic", True)
                elif tag in {"/b", "/strong", "/i"} and content.strip():
                    run = p.add_run(content)
                    setattr(run, "bold" if tag in {"/b", "/strong"} else "italic", False)
                elif tag in {"/h1", "/h2", "/h3", "/h4", "/ol", "/ul"}:
                    style_to_remove = heading_styles.get(tag[1:]) or list_styles.get(tag[1:])
                    if style_to_remove:
                        del style[last_index(style, style_to_remove)]
                        if tag in {"/ol", "/ul"}:
                            list_level -= 1
                j += 1

            output_filename = data.ClientName+"_Retainer_Letter.docx"
            doc.save(data.FilePath+"/"+output_filename)

        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    
    return StreamingResponse(content_generator(), media_type="text/event-stream")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
